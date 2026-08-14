import gc
import json
import re
import warnings
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple, List, Dict
from collections import Counter

import docx
import fitz
import ftfy
import pandas as pd
import pymupdf4llm

# Importación de la configuración centralizada V2
from settings import settings


# ------------------------------------------------------------------------------
# CARGA DE REGLAS DE LIMPIEZA Y RUTAS DESDE SETTINGS
# ------------------------------------------------------------------------------
def cargar_cleaning_rules() -> dict:
    """Carga las reglas de limpieza desde el archivo JSON definido en settings."""
    ruta_rules = settings.CLEANING_RULES_PATH
    if ruta_rules.exists():
        with open(ruta_rules, "r", encoding="utf-8") as f:
            return json.load(f)
    warnings.warn(f"No se encontró el archivo de reglas en: {ruta_rules}")
    return {}


# Reglas de limpieza disponibles globalmente en el módulo
CLEANING_RULES = cargar_cleaning_rules()

# Directorio de entrada/salida tomado directamente de settings
RUTAS_ARCHIVOS = settings.ARCHIVOS_DIR


warnings.filterwarnings("ignore")

# ==============================================================================
# 1. DETECCIÓN Y LIMPIEZA DE ENCABEZADOS / PIES / RUIDO
# ==============================================================================
def detectar_headers_footers_repetidos(
    doc: fitz.Document, min_paginas: int = 3
) -> Tuple[List[str], List[str]]:
    if len(doc) < min_paginas:
        return [], []
    
    indices = sorted(list(set(list(range(min(5, len(doc)))) + list(range(max(0, len(doc)-5), len(doc))))))
    headers, footers = [], []
    
    for idx in indices:
        blocks = sorted(doc[idx].get_text("blocks"), key=lambda b: b[1])
        if not blocks:
            continue
            
        if blocks[0][6] == 0 and 0 < len(blocks[0][4].strip()) < 150:
            headers.append(blocks[0][4].strip())
            
        if len(blocks) > 1 and blocks[-1][6] == 0 and 0 < len(blocks[-1][4].strip()) < 150:
            footers.append(blocks[-1][4].strip())
            
    headers_comunes = [txt for txt, count in Counter([h.lower() for h in headers]).items() if count >= min_paginas]
    footers_comunes = [txt for txt, count in Counter([f.lower() for f in footers]).items() if count >= min_paginas]
    
    return headers_comunes, footers_comunes


def eliminar_ruido_repetido(
    lineas: List[str], 
    tipo_documento: str,
    headers_comunes: List[str],
    footers_comunes: List[str]
) -> List[str]:
    if not lineas:
        return []
    
    frecuencia = Counter()
    patrones_preservar = [re.compile(p, re.I) for p in [r"(bcn\.cl|leychile\.cl|diariooficial\.cl)", r"(D\.O\.?|Diario Oficial)"]]
    
    for linea in lineas:
        norm = linea.strip().lower()
        if norm and len(norm) >= 5:
            if tipo_documento == "LEY_CHILE" and any(p.search(linea) for p in patrones_preservar):
                continue
            frecuencia[norm] += 1
            
    lineas_repetidas = {norm for norm, count in frecuencia.items() if count >= 3}
    lineas_repetidas.update(h.lower().strip() for h in headers_comunes)
    lineas_repetidas.update(f.lower().strip() for f in footers_comunes)
    
    return [l for l in lineas if l.strip().lower() not in lineas_repetidas]


def es_url_oficial(url: str) -> bool:
    dominios = [r"bcn\.cl", r"leychile\.cl", r"diariooficial\.cl", r"gob\.cl", r"sii\.cl", r"contraloria\.cl"]
    return any(re.search(p, url, re.I) for p in dominios)


def limpiar_html_y_enlaces(texto: str) -> str:
    if not texto:
        return texto
    texto = re.sub(r'<a\s+[^>]*href="[^"]*"[^>]*>(.*?)</a>', r"\1", texto, flags=re.I | re.S)
    texto = re.sub(r'<[^>]+>', '', texto)
    
    def reemplazar_md(match):
        return f"[{match.group(1)}]({match.group(2)})" if es_url_oficial(match.group(2)) else match.group(1)
        
    texto = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', reemplazar_md, texto)
    return re.sub(r"[ \t]{2,}", " ", texto)


def eliminar_artefactos(lineas: List[str]) -> List[str]:
    patrones = [re.compile(p, re.I) for p in [r"^'''", r"^```", r"^---+\s*$", r"^\*\*\*+\s*$", r"^===+\s*$"]]
    res = []
    for l in lineas:
        limp = l.strip()
        if any(p.match(limp) for p in patrones) or re.match(r"^[\s'\"\-\=\*\#\~]+$", limp):
            continue
        res.append(l)
    return res

# ==============================================================================
# 2. PROCESAMIENTO TEXTUAL Y ESTRUCTURA
# ==============================================================================
def detectar_tipo_documento(texto: str) -> str:
    muestra = texto[:3000] if texto else ""
    patrones = CLEANING_RULES.get("document_type_patterns", {})
    
    for tipo, env_key in [("LEY_CHILE", "ley_chile"), ("ISO_HSEQ", "iso_hseq"), 
                          ("PRESENTACION", "presentacion"), ("NORMATIVO", "normativo_extenso")]:
        if any(re.search(p, muestra, re.I) for p in patrones.get(env_key, [])):
            return tipo
    return "GENERAL"


def separar_indice_de_contenido(lineas: List[str]) -> Tuple[List[str], List[str]]:
    indice, contenido, en_indice = [], [], False
    patron_simple = re.compile(CLEANING_RULES.get("index_line_patterns", {}).get("linea_indice_simple", r"^$"))
    
    for l in lineas:
        limp = l.strip()
        if re.match(r"^(ÍNDICE|INDICE|CONTENIDO|TABLA\s+DE\s+CONTENIDO)", limp, re.I):
            en_indice = True
            indice.append(limp)
            continue
        if en_indice:
            if patron_simple.match(limp):
                indice.append(limp)
                continue
            en_indice = False
        contenido.append(limp)
        
    return contenido, indice


def es_linea_estructural(linea: str, tipo_doc: str) -> bool:
    limp = linea.strip()
    if not limp:
        return False
        
    if any(re.match(p, limp, re.I) for p in CLEANING_RULES.get("do_not_unify_patterns", [])):
        return True
        
    if tipo_doc in ["ISO_HSEQ", "NORMATIVO"]:
        for p in CLEANING_RULES.get("hierarchical_numbering", {}).values():
            if re.match(p, limp) and len(limp.split()) <= 8:
                return True
                
    return limp == limp.upper() and 3 <= len(limp.split()) <= 10 and not re.search(r"[\.\?\!]$", limp)


def reconstruir_parrafos(lineas: List[str], tipo_doc: str) -> List[str]:
    lineas = eliminar_artefactos(lineas)
    resultado, buffer = [], ""
    preservar_hifen = CLEANING_RULES.get("hyphenation_rules", {}).get("preserve_hyphenated", [])
    
    for l in lineas:
        limp = l.strip()
        if not limp:
            continue
            
        if es_linea_estructural(limp, tipo_doc):
            if buffer:
                resultado.append(buffer)
                buffer = ""
            resultado.append(limp)
            continue
            
        if buffer:
            if buffer.rstrip().endswith("-") and not any(p.lower() in buffer.lower() for p in preservar_hifen):
                buffer = buffer.rstrip()[:-1] + limp
            elif re.search(r"[\;\,\:\.]\s*$", buffer):
                resultado.append(buffer)
                buffer = limp
            else:
                buffer = f"{buffer} {limp}"
        else:
            buffer = limp
            
    if buffer:
        resultado.append(buffer)
    return resultado


def detectar_y_formatear_titulos(lineas: List[str], tipo_doc: str) -> List[str]:
    resultado = []
    patrones_titulos = CLEANING_RULES.get("structural_titles", {})
    
    patrones_no_titulo = [
        re.compile(p, re.I) for p in [
            r"^'''", r"^```", r"^D\.O\.?", r"^D\.F\.L\.?", r"^N[°º.º]?\s*\d+", 
            r"^\d{2}\.\d{2}\.\d{4}", r"^(Fecha|Tipo|Inicio|Fin|Ley)\s+", 
            r"^www\.", r"^Biblioteca\s+del\s+Congreso", r"^documento\s+generado",
            r"^página\s+\d+\s+de\s+\d+", r"^NOTA\s*$", r"^Visto\s*:"
        ]
    ]
    
    for l in lineas:
        limp = l.strip()
        if not limp or re.match(r"^#{1,6}\s+", limp) or any(p.match(limp) for p in patrones_no_titulo) or re.match(r"^https?://", limp):
            resultado.append(limp)
            continue
            
        es_titulo, nivel = False, 3
        
        for niv_cfg, kw_list in patrones_titulos.items():
            if any(re.match(rf"^{kw}\b.+$", limp, re.I) for kw in kw_list):
                es_titulo, nivel = True, (2 if niv_cfg == "level_2" else 3)
                break
                
        if not es_titulo and re.match(r"^Artículo\s+\d+", limp, re.I):
            es_titulo, nivel = True, 2
        elif not es_titulo and re.match(r"^(Título|Capítulo|Párrafo|Sección)\s+[IVX]+", limp, re.I):
            es_titulo, nivel = True, 2
        elif not es_titulo and tipo_doc in ["ISO_HSEQ", "NORMATIVO"]:
            m = re.match(r"^(\d+\.)*\d+[\.\)]\s+(.+)$", limp)
            if m and len(m.group(2).strip()) > 2:
                es_titulo, nivel = True, min(m.group(1).count(".") + 1, 6)
                
        resultado.append(f"{'#' * nivel} {limp}" if es_titulo else limp)
        
    return resultado


def formatear_listas(lineas: List[str]) -> List[str]:
    res = []
    for l in lineas:
        limp = l.strip()
        if re.match(r"^[•\-\*\+]\s+", limp):
            res.append(f"- {limp[1:].strip()}")
        elif re.match(r"^([a-z]|\d+)\)\s+", limp):
            res.append(f"- {limp}")
        else:
            res.append(limp)
    return res

# ==============================================================================
# 3. EXTRACTION ENGINES
# ==============================================================================
def procesar_pdf_robusto(ruta_pdf: Path) -> str:
    try:
        doc = fitz.open(str(ruta_pdf))
        headers_c, footers_c = detectar_headers_footers_repetidos(doc)
        paginas_md = pymupdf4llm.to_markdown(str(ruta_pdf), page_chunks=True)
        doc.close()
        
        texto_crudo = "\n\n".join([p.get("text", "").strip() for p in paginas_md if p.get("text")])
        if not texto_crudo:
            return "[Error: No se pudo extraer texto del PDF]"
            
        texto = limpiar_html_y_enlaces(ftfy.fix_text(texto_crudo))
        lineas = [l.strip() for l in texto.split("\n") if l.strip()]
        
        tipo_doc = detectar_tipo_documento(texto)
        lineas = eliminar_ruido_repetido(lineas, tipo_doc, headers_c, footers_c)
        lineas, _ = separar_indice_de_contenido(lineas)
        lineas = reconstruir_parrafos(lineas, tipo_doc)
        lineas = detectar_y_formatear_titulos(lineas, tipo_doc)
        lineas = formatear_listas(lineas)
        
        return re.sub(r"\n{3,}", "\n\n", "\n\n".join(lineas)).strip()
    except Exception as e:
        return f"[Error procesando PDF: {str(e)}]"


def leer_docx(ruta: Path) -> str:
    try:
        doc = docx.Document(ruta)
        res = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            estilo = (p.style.name or "").lower() if p.style else ""
            m = re.match(r"heading\s+(\d+)", estilo)
            res.append(f"{'#' * min(int(m.group(1)), 6)} {txt}" if m else (f"# {txt}" if estilo == "title" else txt))
            
        for tabla in doc.tables:
            res.append("")
            for fila in tabla.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if celdas:
                    res.append("| " + " | ".join(celdas) + " |")
        return "\n\n".join(res)
    except Exception as e:
        return f"[Error leyendo DOCX: {str(e)}]"


def leer_excel(ruta: Path) -> str:
    try:
        hojas = pd.read_excel(ruta, sheet_name=None)
        return "\n\n".join([f"## Hoja: {n}\n" + df.to_markdown(index=False) for n, df in hojas.items()])
    except Exception as e:
        return f"[Error leyendo Excel: {str(e)}]"

# ==============================================================================
# 4. PIPELINE DE EJECUCIÓN
# ==============================================================================
def procesar_archivo(ruta_in: Path, ruta_out: Path, categoria: str) -> bool:
    if ruta_out.exists():
        return False
        
    ext = ruta_in.suffix.lower()
    if ext == ".pdf":
        contenido = procesar_pdf_robusto(ruta_in)
    elif ext == ".docx":
        contenido = leer_docx(ruta_in)
    elif ext == ".txt":
        contenido = ruta_in.read_text(encoding="utf-8", errors="ignore")
    elif ext in [".xlsx", ".xls"]:
        contenido = leer_excel(ruta_in)
    else:
        return False

    frontmatter = (
        f"---\ntitle: \"{ruta_in.stem}\"\nsource_file: \"{ruta_in.name}\"\n"
        f"category: \"{categoria}\"\nprocessed_date: \"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\"\n---\n\n"
        f"# {ruta_in.stem}\n\n"
    )
    ruta_out.write_text(frontmatter + contenido, encoding="utf-8")
    return True


def ejecutar_pipeline():
    print(f"[{datetime.now().strftime('%H:%M:%S')}] 🚀 Iniciando pipeline...")
    tot_proc, tot_omit = 0, 0
    exts = {".pdf", ".docx", ".txt", ".xlsx", ".xls"}
    
    for cfg in settings.RUTAS_CATEGORIAS:
        cat = cfg["categoria"]
        p_in, p_out = cfg["input_dir"], cfg["output_dir"]
        
        if not p_in.exists():
            print(f"⚠️ Omite categoría {cat}: No existe el directorio {p_in}")
            continue
            
        p_out.mkdir(parents=True, exist_ok=True)
        archivos = [f for f in p_in.rglob("*") if f.is_file() and f.suffix.lower() in exts]
        
        proc, omit = 0, 0
        for i, arc in enumerate(archivos, 1):
            if procesar_archivo(arc, p_out / f"{arc.stem}.md", cat):
                proc += 1
            else:
                omit += 1
            if i % 10 == 0:
                gc.collect()
                
        tot_proc += proc
        tot_omit += omit
        print(f"✔ {cat} -> Nuevos: {proc} | Omitidos: {omit}")

    print(f"🎉 Finalizado -> Procesados: {tot_proc} | Omitidos: {tot_omit}")


if __name__ == "__main__":
    ejecutar_pipeline()