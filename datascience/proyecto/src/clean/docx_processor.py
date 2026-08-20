"""Extracción de texto desde DOCX, priorizando 'unstructured' con fallback a
python-docx."""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List

import docx

from .modelos import LineaFuente

logger = logging.getLogger("clean.docx_processor")

try:
    from unstructured.partition.docx import partition_docx
    UNSTRUCTURED_DISPONIBLE = True
except ImportError:
    UNSTRUCTURED_DISPONIBLE = False

TIPOS_RUIDO_UNSTRUCTURED = {"Header", "Footer", "PageBreak"}


def _extraer_con_unstructured(ruta: Path) -> List[LineaFuente]:
    elementos = partition_docx(filename=str(ruta))
    lineas = []
    for i, el in enumerate(elementos, start=1):
        if type(el).__name__ in TIPOS_RUIDO_UNSTRUCTURED:
            continue
        texto = str(el).strip()
        if texto:
            # No se reescribe el texto (p.ej. anteponer "##" a los Title):
            # la detección de jerarquía (estructura.py) se aplica de forma
            # uniforme sobre el texto crudo, sea cual sea el extractor usado.
            lineas.append(LineaFuente(texto=texto, linea_original=i))
    return lineas


def _extraer_con_python_docx(ruta: Path) -> List[LineaFuente]:
    documento = docx.Document(str(ruta))
    lineas = []
    contador = 1

    for p in documento.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        lineas.append(LineaFuente(texto=txt, linea_original=contador))
        contador += 1

    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                lineas.append(LineaFuente(texto="| " + " | ".join(celdas) + " |", linea_original=contador))
                contador += 1

    return lineas


def extraer_lineas_docx(ruta_docx: Path, forzar_fallback: bool = False) -> List[LineaFuente]:
    if UNSTRUCTURED_DISPONIBLE and not forzar_fallback:
        try:
            return _extraer_con_unstructured(ruta_docx)
        except Exception as e:
            logger.warning("unstructured falló en %s (%s); usando fallback python-docx.", ruta_docx.name, e)
    return _extraer_con_python_docx(ruta_docx)
